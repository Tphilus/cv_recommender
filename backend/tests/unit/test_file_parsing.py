import io

import pytest
from docx import Document
from pypdf import PdfWriter

from app.utils.file_parsing import extract_docx_text, extract_pdf_text, extract_plain_text, to_image_bytes


def test_to_image_bytes_passes_through_image_unchanged():
    raw = b"fake-png-bytes"
    out_bytes, out_mime = to_image_bytes(raw, "image/png")
    assert out_bytes == raw
    assert out_mime == "image/png"


def test_to_image_bytes_rejects_unsupported_mime():
    with pytest.raises(ValueError):
        to_image_bytes(b"data", "application/zip")


def test_extract_docx_text_reads_paragraphs():
    doc = Document()
    doc.add_paragraph("Ama Owusu")
    doc.add_paragraph("Backend Engineer")
    buf = io.BytesIO()
    doc.save(buf)

    text = extract_docx_text(buf.getvalue())

    assert "Ama Owusu" in text
    assert "Backend Engineer" in text


def test_extract_pdf_text_reads_page_text(monkeypatch):
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buf = io.BytesIO()
    writer.write(buf)

    # A blank page has no real text; patch extract_text to simulate a real CV page
    # rather than depending on pypdf's ability to draw+extract text in a unit test.
    from pypdf import PageObject

    monkeypatch.setattr(PageObject, "extract_text", lambda self: "Ama Owusu\nBackend Engineer")

    text = extract_pdf_text(buf.getvalue())

    assert "Ama Owusu" in text


def test_extract_pdf_text_returns_empty_string_on_scanned_pdf():
    writer = PdfWriter()
    writer.add_blank_page(width=200, height=200)
    buf = io.BytesIO()
    writer.write(buf)

    # No text layer (scanned/image-only PDF) — callers use the empty string to
    # decide on a fallback strategy (e.g. pdf_to_image_bytes) rather than a raised
    # exception, since this is an expected, recoverable case, not a real error.
    assert extract_pdf_text(buf.getvalue()) == ""


def test_extract_plain_text_decodes_utf8():
    text = extract_plain_text("Ama Owusu\nBackend Engineer".encode("utf-8"))
    assert "Ama Owusu" in text
